# import docx NOT python-docx
import docx
import re



# create an instance of a word document
doc = docx.Document()



###input header
print("Geben Sie die Überschrift ein: ")
headline = input()
doc.add_heading(headline, 0)


### input date and add to word doc
print("Gebene Sie das Datum ein:")
date = input()
doc.add_paragraph(date)

print("Geben Sie die Klasse ein:")
klasse = input()
doc.add_paragraph(klasse)


# turn the headline into a viable name for the file.
filename = headline + date + ".docx"

filename = re.sub(" ", "_", filename)

filename = re.sub("/", "_", filename)


###input text
print("Und nun zum Text")
print("Geben Sie hier einleitende Worte zum Text ein (optional):")
info = input()
doc.add_paragraph(info)


print("Geben Sie den zu übersetzenden Text ein:")
text = input()
sätze = text.split(". ")

###inputaufgaben

def aufgaben():
    print("Geben Sie eine Aufgabe ein:")

def aufgabendrucken():
    for i in sätze:
        print(i, ".")

#output
print(headline)
print(date)
print(klasse)
print(info)
aufgabendrucken()


  
# add a heading of level 0 (largest heading)
#doc.add_heading(headline, 0)
  
# add a paragraph and store 
# the object in a variable

#doc.add_paragraph(date)
  
# add a run i.e, style like 
# bold, italic, underline, etc.
#doc_para.add_run('a very important point').bold = True
#doc_para.add_run(' that needs to be accompanied by ')
#doc_para.add_run('emphasis').italic = True

  
# now save the document to a location
doc.save(filename)